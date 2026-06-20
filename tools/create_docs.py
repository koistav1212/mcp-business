import os
from typing import Optional, Type
from pydantic import BaseModel
from tools.base import BaseTool
from storage.local import local_storage
from docx import Document
from docx.shared import Pt, RGBColor

class CreateDocsInput(BaseModel):
    filename: str
    title: str
    body_text: str
    style: Optional[str] = "professional"

class CreateDocsTool(BaseTool):
    name: str = "create_docs"
    description: str = "Generates a beautifully styled Microsoft Word (.docx) document from text paragraphs."
    args_schema: Optional[Type[BaseModel]] = CreateDocsInput

    async def execute(self, **kwargs) -> str:
        filename = kwargs["filename"]
        if not filename.endswith(".docx"):
            filename += ".docx"
            
        title_text = kwargs["title"]
        body_text = kwargs["body_text"]
        style = kwargs.get("style", "professional")

        doc = Document()
        
        # Color mapping based on style
        if style == "creative":
            primary_color = RGBColor(99, 102, 241)  # Indigo
            text_color = RGBColor(79, 70, 229)     # Violet
        elif style == "minimalist":
            primary_color = RGBColor(9, 9, 11)      # Black
            text_color = RGBColor(39, 39, 42)      # Charcoal
        else:  # professional
            primary_color = RGBColor(30, 41, 59)    # Dark Slate
            text_color = RGBColor(51, 65, 85)      # Slate Grey

        # Add Title with styling
        title_p = doc.add_paragraph()
        run = title_p.add_run(title_text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = primary_color
        
        # Add Spacing
        title_p.paragraph_format.space_after = Pt(20)

        # Split and add body paragraphs
        paragraphs = body_text.split("\n\n")
        for p in paragraphs:
            cleaned_p = p.strip()
            if cleaned_p:
                para = doc.add_paragraph()
                # Split inside paragraph to handle single newlines as line breaks
                lines = cleaned_p.split("\n")
                for index, line in enumerate(lines):
                    if index > 0:
                        para.add_run("\n")
                    run = para.add_run(line.strip())
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.color.rgb = text_color
                para.paragraph_format.space_after = Pt(10)

        # Save to temporary path
        temp_path = f"temp_gen_{filename}"
        doc.save(temp_path)
        
        # Read bytes and upload to storage
        with open(temp_path, "rb") as f:
            content = f.read()
            
        if os.path.exists(temp_path):
            os.remove(temp_path)

        download_url = local_storage.write_content(filename, content, "docs")
        return download_url
